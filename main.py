# streamlit_app.py

import streamlit as st
import fitz  # PyMuPDF for PDFs
import docx  # for .docx files
import os
from dotenv import load_dotenv

from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.schema import Document

from langchain_huggingface import HuggingFaceEndpoint, ChatHuggingFace
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser

# === Load .env ===
load_dotenv()
hf_token = os.getenv("HUGGINGFACEHUB_API_TOKEN")

# === UI Header ===
st.title("📄 Document-Based Q&A System")

# === ⚠️ Warning for .env setup ===
st.warning(
    "⚠️ Please ensure you have a `.env` folder with your API key set as Required Format.\n"
    "You may replace the model with any supported LLM by changing 'Minimal' code."
)

# === 🔽 Dropdown to Explain Techniques ===
with st.expander("🔍 What Technique is Used?"):
    technique = st.selectbox("Select a component to learn about it:", [
        "Document Chunking",
        "Vector Embedding with Hugging Face",
        "FAISS Vector Store for Similarity Search",
        "LLM-based Contextual Answer Generation"
    ])

    explanations = {
        "Document Chunking": "The uploaded document is split into overlapping text chunks using RecursiveCharacterTextSplitter.",
        "Vector Embedding with Hugging Face": "Each chunk is converted into a numerical vector using sentence-transformers from Hugging Face.",
        "FAISS Vector Store for Similarity Search": "The embedded vectors are stored in a FAISS index for fast similarity search based on user queries.",
        "LLM-based Contextual Answer Generation": "The most relevant chunks are passed to a Hugging Face LLM to generate an answer."
    }

    st.info(explanations[technique])

# === Load Hugging Face LLM ===
llm = HuggingFaceEndpoint(
    repo_id="mistralai/Mixtral-8x7B-Instruct-v0.1",
    task="text-generation",
    huggingfacehub_api_token=hf_token
)
model = ChatHuggingFace(llm=llm)
parser = StrOutputParser()

# === Prompt Template ===
qa_prompt = PromptTemplate(
    template="Use the following context to answer the question.\n\nContext:\n{context}\n\nQuestion:\n{question}",
    input_variables=["context", "question"]
)

# === Embedding Model ===
embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# === File Upload ===
uploaded_file = st.file_uploader("📎 Upload a PDF or Word document", type=["pdf", "docx"])

# === Text Extraction Functions ===
def extract_text_from_pdf(file):
    text = ""
    with fitz.open(stream=file.read(), filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text.strip()

def extract_text_from_docx(file):
    text = ""
    doc = docx.Document(file)
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text.strip()

# === Process File Upload ===
if uploaded_file:
    if uploaded_file.name.endswith(".pdf"):
        raw_text = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.name.endswith(".docx"):
        raw_text = extract_text_from_docx(uploaded_file)
    else:
        st.error("❌ Unsupported file format.")
        st.stop()

    # === Chunking the Text ===
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    chunks = text_splitter.split_text(raw_text)
    documents = [Document(page_content=chunk) for chunk in chunks]

    # === Embedding and Vector Indexing ===
    with st.spinner("🔍 Indexing document with FAISS..."):
        vectorstore = FAISS.from_documents(documents, embedding_model)
    st.success("✅ Document processed and indexed!")

    # === Q&A Interface ===
    st.subheader("❓ Ask a Question About the Document")
    user_question = st.text_input("Type your question here:")

    if user_question:
        with st.spinner("💭 Thinking..."):
            # Step 1: Similarity Search
            similar_docs = vectorstore.similarity_search(user_question, k=3)
            context = "\n\n".join([doc.page_content for doc in similar_docs])

            # Step 2: Generate Answer with Context
            chain = qa_prompt | model | parser
            answer = chain.invoke({"context": context, "question": user_question})

        st.markdown("**Answer:**")
        st.write(answer)

# === Footer ===
st.markdown("---")
st.markdown("Made with ❤️ by **Shivendra** using Hugging Face, FAISS, and Streamlit")
st.markdown("© 2025 Shivendra Prasad Mishra. All rights reserved.")
